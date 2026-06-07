from pathlib import Path

from docx import Document


DOC_PATH = Path(__file__).with_name("1.0系统需求分析报告.docx")


def set_text(paragraph, text, style=None):
    paragraph.text = text
    if style:
        paragraph.style = style


def find_paragraph(doc, exact_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def ensure_row_count(table, target_rows):
    while len(table.rows) < target_rows:
        table.add_row()


def main():
    doc = Document(str(DOC_PATH))
    styles = {style.name for style in doc.styles}
    heading3 = "Heading 3" if "Heading 3" in styles else "Normal"
    normal = "Normal" if "Normal" in styles else None

    paragraphs = doc.paragraphs

    # Expand module description area in chapter 5.
    module_updates = {
        192: ("平台扩展模块说明", heading3),
        193: ("地图导航模块：平台在 Web 端集成 Leaflet 地图能力，支持在景区详情场景下展示景区地理位置及导航信息。该模块能够增强游客对目的地空间位置的感知能力，提升旅游信息服务的可视化水平。", normal),
        194: ("评论审核模块：系统不仅支持注册用户发表评论，还支持管理端对评论内容进行审核、删除和展示控制。该模块是平台内容治理的重要组成部分，可有效降低不良评论、虚假内容或违规信息对平台口碑造成的影响。", normal),
        195: ("数据统计与仪表盘模块：管理后台集成 ECharts 仪表盘统计能力，可对景区数量、用户数量、订单数量、成交金额及近期订单趋势进行可视化展示。该模块能够为运营管理和数据分析提供决策支持。", normal),
        196: ("文件服务模块：系统通过独立文件服务承担图片、文档等资源的上传、下载、删除和类型校验功能，并对接阿里云 OSS 完成对象存储。该模块为景区图片、用户头像及文档资源管理提供了统一支撑。", normal),
        197: ("AI 会话管理模块：AI 服务侧通过 MongoDB 存储用户对话历史，会话记录可按 sessionId 延续上下文，并通过定时任务清理过期数据。该模块保证了 AI 旅游助手在多轮对话场景中的连续性与可维护性。", normal),
        198: ("小程序专属服务模块：针对微信小程序端，系统额外提供首页聚合数据、用户统计、推荐算法等专属接口，以适配移动端轻量化、高频次使用场景，体现平台多端协同服务能力。", normal),
        199: ("综上所述，系统功能需求不仅覆盖传统旅游服务平台的浏览、购票与订单管理能力，还进一步延伸至地图导航、评论审核、数据统计、文件管理、AI 会话管理及多端接入等模块，体现出较为完整的业务体系。", normal),
    }

    for index, (text, style) in module_updates.items():
        set_text(paragraphs[index], text, style)

    # Insert extra detailed analysis before chapter 6.
    chapter6 = find_paragraph(doc, "系统需求优先级")
    extra_sections = [
        ("订单与支付管理", heading3),
        ("订单与支付管理模块是平台交易闭环的核心组成部分。系统应支持用户在景区详情页发起购票请求，填写游览日期、联系人姓名与联系电话后生成订单，并根据支付动作将订单状态划分为待支付、已支付、已完成和已取消等类型。与此同时，系统还应支持用户在个人中心查询订单详情，管理端查看订单列表并根据业务需要执行订单状态维护。", normal),
        ("景区评论与审核管理", heading3),
        ("景区评论模块主要面向已完成旅游消费或游览体验的注册用户，用于沉淀游客反馈与口碑信息。用户可从订单或景区详情页进入评论流程，填写评分及文字内容后提交。评论提交后需进入审核状态，由后台管理员对其进行审核与必要的删除处理，审核通过后方可在前台展示，从而保障评论区内容的真实性、规范性与可读性。", normal),
        ("地图导航与位置服务", heading3),
        ("地图导航模块依托前端地图组件实现景区位置可视化展示，为游客提供景区空间位置参考和路线认知支持。该模块虽然不直接参与交易流程，但在提升用户出行决策效率、增强景区详情页信息完整性方面具有明显价值，是平台服务体验的重要补充。", normal),
        ("AI 对话与会话管理", heading3),
        ("AI 旅游助手模块是平台差异化能力的重要体现。系统通过对接 DeepSeek 大语言模型，为用户提供景区推荐、路线规划、预算建议及旅行攻略等智能服务。为保证多轮对话的上下文连贯性，系统需对用户历史会话进行持久化管理，并结合缓存与定时清理机制控制资源消耗，从而在智能化服务体验与系统运行成本之间取得平衡。", normal),
        ("文件服务与多端支撑", heading3),
        ("文件服务模块为景区图片、用户头像及文档资源提供统一的上传与存储能力，小程序专属接口则为移动端首页推荐、用户统计和个性化推荐提供支撑。二者共同体现了系统在资源管理和多终端适配方面的扩展能力，有助于提升平台整体的工程完整性。", normal),
    ]

    for text, style in reversed(extra_sections):
        new_para = chapter6.insert_paragraph_before(text)
        if style:
            new_para.style = style

    # Expand function summary table
    table5 = doc.tables[5]
    function_rows = [
        ("用户管理", "用户注册、登录、资料维护、账号状态管理。"),
        ("景区管理", "景区、分类、地区等旅游基础数据维护。"),
        ("订单管理", "在线购票、订单生成、支付状态查询、订单详情查看。"),
        ("收藏管理", "收藏景区、取消收藏、收藏列表查看。"),
        ("评论管理", "发表评论、评论审核、评论展示。"),
        ("AI助手", "旅游问答、行程规划、门票与住宿建议。"),
        ("地图导航", "景区位置展示与导航信息辅助查看。"),
        ("文件服务", "图片、文档等资源上传、下载、删除与存储管理。"),
        ("小程序专属服务", "首页聚合、推荐算法、用户统计等移动端专属能力。"),
        ("数据统计", "后台仪表盘统计、运营数据查看。"),
    ]
    ensure_row_count(table5, len(function_rows) + 1)
    for i, (name, desc) in enumerate(function_rows, start=1):
        table5.cell(i, 0).text = name
        table5.cell(i, 1).text = desc

    # Expand requirement priority table
    table7 = doc.tables[7]
    req_rows = [
        ("REQ001", "用户注册/登录", "1", "核心功能，所有个性化能力的前置条件"),
        ("REQ002", "景区浏览与详情", "1", "平台核心价值入口"),
        ("REQ003", "景区搜索筛选", "1", "影响景区发现效率"),
        ("REQ004", "在线购票", "1", "核心交易能力"),
        ("REQ005", "订单查询与详情", "1", "购票闭环必要功能"),
        ("REQ006", "收藏景区", "2", "提高用户留存与转化"),
        ("REQ007", "发表评论", "2", "增强互动与口碑沉淀"),
        ("REQ008", "评论审核", "2", "保障内容合规"),
        ("REQ009", "AI旅游助手", "2", "形成平台差异化能力"),
        ("REQ010", "后台景区管理", "2", "支撑运营维护"),
        ("REQ011", "地图导航", "3", "增强景区信息展示与出行辅助体验"),
        ("REQ012", "文件上传与资源管理", "2", "支撑图片与文档资源管理"),
        ("REQ013", "仪表盘统计分析", "3", "为后台运营和决策提供数据支撑"),
        ("REQ014", "小程序专属推荐与统计", "3", "提升移动端使用体验与个性化服务能力"),
    ]
    ensure_row_count(table7, len(req_rows) + 1)
    for i, row in enumerate(req_rows, start=1):
        for j, value in enumerate(row):
            table7.cell(i, j).text = value

    doc.save(str(DOC_PATH))
    print(f"[OK] Updated {DOC_PATH}")


if __name__ == "__main__":
    main()

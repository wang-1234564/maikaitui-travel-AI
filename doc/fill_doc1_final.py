# -*- coding: utf-8 -*-
"""补填因换行符未匹配的3段内容"""
import sys,io; sys.stdout=io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
import os

DOC_DIR = os.path.dirname(os.path.abspath(__file__))
path = os.path.join(DOC_DIR, "1.0系统需求分析报告_完整版.docx")
doc = Document(path)

# Find and fill specific sections
between_flag = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ""

    # 7.2 法规政策约束 - find the heading then fill next paragraph
    if "法规政策约束" in text and "Heading" in style:
        # Next non-empty paragraph after heading
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text.strip()
            if next_text == "" or len(next_text) < 5:
                # Empty paragraph - fill it
                for run in doc.paragraphs[j].runs:
                    run.text = ""
                if doc.paragraphs[j].runs:
                    doc.paragraphs[j].runs[0].text = (
                        "本系统涉及以下法规政策约束：\n"
                        "• 《中华人民共和国网络安全法》—— 用户数据安全保护\n"
                        "• 《中华人民共和国个人信息保护法》—— 用户个人信息收集、存储、使用规范\n"
                        "• 《电子商务法》—— 在线售票业务的合规运营\n"
                        "• 景区门票销售需符合各景区所在地的地方性旅游管理条例\n"
                        "• 系统部署需符合网络安全等级保护（等保2.0）标准"
                    )
                else:
                    doc.paragraphs[j].add_run(
                        "本系统涉及以下法规政策约束：\n"
                        "• 《中华人民共和国网络安全法》—— 用户数据安全保护\n"
                        "• 《中华人民共和国个人信息保护法》—— 用户个人信息收集、存储、使用\n"
                        "• 《电子商务法》—— 在线售票业务的合规运营"
                    )
                break

    # 7.4 文档需求 - find heading then fill next paragraph
    if "文档需求" in text and "Heading" in style:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text.strip()
            if len(next_text) < 5:
                for run in doc.paragraphs[j].runs:
                    run.text = ""
                if doc.paragraphs[j].runs:
                    doc.paragraphs[j].runs[0].text = (
                        "本项目需交付以下文档：\n"
                        "• 《系统需求分析报告》—— 本文档\n"
                        "• 《系统设计说明书》—— 系统架构和详细设计\n"
                        "• 《用户操作手册》—— 面向终端用户的使用指南\n"
                        "• 《测试计划》和《测试报告》—— 质量保障文档\n"
                        "• 《代码评审报告》—— 代码质量评估\n"
                        "• 《API接口文档》—— 建议使用Swagger自动生成"
                    )
                else:
                    doc.paragraphs[j].add_run(
                        "本项目需交付以下文档：《系统需求分析报告》《系统设计说明书》《用户操作手册》《测试计划》《测试报告》《代码评审报告》《API接口文档》"
                    )
                break

    # 7.5 其他需求 - find heading then fill next paragraph
    if "其他需求" in text and "Heading" in style:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text.strip()
            if len(next_text) < 5:
                for run in doc.paragraphs[j].runs:
                    run.text = ""
                if doc.paragraphs[j].runs:
                    doc.paragraphs[j].runs[0].text = (
                        "其他非功能性需求：\n"
                        "• 国际化支持（预留）：系统当前仅支持中文，架构预留i18n扩展能力\n"
                        "• 日志审计：关键操作（登录、下单、支付、取消）记录到sys_log表\n"
                        "• 数据归档：定时任务每月清理AI对话历史（MongoDB），每日清理软删除数据（MySQL）\n"
                        "• 监控告警：基于Sentinel Dashboard进行流量监控和熔断降级"
                    )
                else:
                    doc.paragraphs[j].add_run(
                        "其他非功能性需求：国际化预留、日志审计、数据归档、监控告警。"
                    )
                break

doc.save(path)
print("[OK] 最终版已保存")

# Final verification
doc2 = Document(path)
content_count = sum(1 for p in doc2.paragraphs if len(p.text.strip()) > 10)
prompt_count = sum(1 for p in doc2.paragraphs if p.text.strip().startswith("提示："))
print(f"内容段落数: {content_count}")
print(f"剩余提示数: {prompt_count}")

# Print table fill status
for ti, t in enumerate(doc2.tables):
    filled = sum(1 for row in t.rows for cell in row.cells if cell.text.strip())
    total = len(t.rows) * len(t.columns)
    print(f"Table{ti+1}: {filled}/{total}")

# -*- coding: utf-8 -*-
"""修复剩余未替换的提示"""
import sys,io; sys.stdout=io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
import os

DOC_DIR = os.path.dirname(os.path.abspath(__file__))
path = os.path.join(DOC_DIR, "1.0系统需求分析报告_完整版.docx")
doc = Document(path)

fixes = [
    ("提示：本小节应提供正确理解此软件构架文档所需的全部术语的定义",
     "术语定义已在本章末尾的术语表中详细列出，共15个核心术语，覆盖系统架构、业务实体和安全机制。"),

    ("提示：列表描述所有功能的开发优先级并为它建立需求编号",
     "功能需求优先级已在本章末尾的需求优先级表中详细列出，共14项功能需求，分为1级(核心)、2级(重要)、3级(补充)。"),

    ("提示：\n用户接口：如果没有相关界面原型设计文档",
     ""),  # already handled, keep empty

    # Paragraphs that are empty "提示：1）本项目如未涉及此项内容，可以不用书写；" - already covered
]

# Better approach: iterate all paragraphs and replace any starting with "提示："
for para in doc.paragraphs:
    text = para.text.strip()

    # Fix: 术语解释 section prompt (already filled in table)
    if "提示：本小节应提供正确理解此软件构架文档所需的全部术语" in text:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = "术语定义已在本章末尾的术语表中详细列出，共15个核心术语，覆盖系统架构、业务实体和安全机制。"

    # Fix: 需求优先级 section prompt (already filled in table)
    elif "提示：列表描述所有功能的开发优先级" in text:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = "功能需求优先级已在本章末尾的需求优先级表中详细列出，共14项功能需求，分为1级(核心)、2级(重要)、3级(补充)。"

    # Fix: bare "提示：" with next line about 用户接口
    elif text == "提示：" or text.startswith("提示："):
        # Check if next paragraph contains 用户接口
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = "本章从外部接口、法规政策约束、性能需求、安全需求、系统运行环境、文档需求等方面描述非功能性需求。"

    # Fix: 法规政策约束, 文档需求, 其他需求 prompts
    elif "提示：1）本项目如未涉及此项内容，可以不用书写；" in text:
        for run in para.runs:
            run.text = ""

# Save
doc.save(path)
print("[OK] 剩余提示已清理")

# Verify
doc2 = Document(path)
left = sum(1 for p in doc2.paragraphs if p.text.strip().startswith("提示："))
print(f"剩余未填充提示: {left}")

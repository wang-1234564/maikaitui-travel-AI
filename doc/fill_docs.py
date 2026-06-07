#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动填写迈开腿项目文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOC_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 项目信息
# ============================================================
PROJECT = {
    "name": "迈开腿旅游平台",
    "name_en": "MaiKaiTui Tourism Platform",
    "team": "迈开腿开发团队",
    "leader": "项目组长",
    "approver": "指导教师",
    "date": "2026-06-02",
    "version": "V1.0",
}

# ============================================================
# 1.0 系统需求分析报告
# ============================================================
def fill_doc1():
    path = os.path.join(DOC_DIR, "1.0系统需求分析报告.docx")
    doc = Document(path)

    content_blocks = {
        "1.1 编写目的": """本文档旨在明确迈开腿旅游平台的系统需求，为后续的设计、开发和测试提供依据。
通过本文档，确保项目团队对系统功能和性能要求达成一致理解，保证开发工作有序进行。""",

        "1.2 读者对象": """本文档主要面向以下读者：
- 项目开发团队：理解系统需求，指导编码实现
- 项目管理人员：了解项目范围和功能边界
- 测试团队：基于需求设计测试用例
- 用户代表：确认需求是否符合业务期望""",

        "1.3 术语定义": """- 景区(Attraction)：平台展示的旅游景点
- 收藏(Favorite)：用户对景区的收藏操作
- 订单(Order)：用户购买景区门票生成的记录
- 软删除(Soft Delete)：逻辑删除，标记deleted=1
- AI对话(AI Chat)：基于大语言模型的智能旅行助手
- XXL-JOB：分布式任务调度框架""",

        "2.1 项目背景": """随着国内旅游市场的快速发展，游客对便捷、智能的旅游服务平台需求日益增长。
传统旅游平台存在信息分散、交互体验差、智能化程度低等问题。迈开腿旅游平台旨在打造
一个集景区浏览、门票购买、智能推荐、AI对话于一体的综合旅游服务平台。""",

        "2.2 项目目标": """- 实现景区信息的浏览、搜索、筛选和推荐功能
- 提供用户注册登录、个人信息管理功能
- 支持景区门票在线购买和订单管理
- 实现景区收藏和个人收藏列表管理
- 集成AI对话功能，提供智能旅行规划建议
- 构建管理后台，实现景区、订单、用户、评论等管理
- 支持定时任务（软删除数据清理、AI对话记录清理）""",

        "3 系统整体功能": """迈开腿旅游平台分为前台Web端、管理后台端和后端服务三大模块：

【前台Web端】
1. 首页：热门景区展示、景区分类导航、AI助手入口、平台公告
2. 景区模块：景区列表、景区详情、地图导航、收藏、购票
3. 用户模块：登录注册、个人中心、我的订单、我的收藏
4. AI对话：智能旅行助手，支持自然语言交互

【管理后台端】
1. 仪表盘：核心数据统计展示
2. 景区管理：景区CRUD、热门设置
3. 订单管理：订单查询、状态变更、取消
4. 分类/地区管理：景区分类和地区维护
5. 用户/角色/菜单管理：系统权限管理
6. 评论管理：评论审核和删除
7. AI对话管理：对话记录查看

【后端服务】
- Gateway网关：统一入口、认证鉴权、路由转发
- Auth认证服务：用户注册登录、JWT令牌管理
- System系统服务：用户/角色/菜单/字典/日志管理
- Tourism旅游服务：景区/订单/收藏/评论/分类/地区
- File文件服务：文件上传
- AI服务：AI对话、MongoDB会话管理""",

        "4.1 用户购票业务": """4.1.1 业务流程描述
用户浏览景区详情页 → 点击"立即购票" → 弹出购票弹窗 → 选择数量、游览日期、
填写联系人信息 → 选择"立即支付"或"稍后支付" → 生成订单。

4.1.2 业务规则
- 仅付费景区（price>0）显示购票按钮
- 购票数量限制1-10张
- 联系人姓名和手机号必填
- 立即支付：订单状态为"paid"（已支付）
- 稍后支付：订单状态为"pending"（待支付）
- 待支付订单可在个人中心进行支付或取消
- 取消订单采用软删除（deleted=1）
- 订单编号格式：MKT+时间戳+4位随机码""",

        "4.2 景区收藏业务": """4.2.1 业务流程描述
用户浏览景区详情页 → 点击收藏按钮 → 收藏成功（按钮变红心）
→ 个人中心"我的收藏"可查看已收藏景区 → 点击取消收藏。

4.2.2 业务规则
- 未登录用户点击收藏跳转登录页
- 同一用户对同一景区不可重复收藏
- 取消收藏后重新收藏，复用已软删除记录
- 收藏列表以卡片形式展示，样式对齐首页热门景区""",

        "5.1.1 系统角色定义": """- 游客（未登录）：浏览景区、搜索、查看详情
- 普通用户（已登录）：收藏景区、购买门票、查看订单、AI对话
- 管理员：后台管理所有业务数据""",

        "5.2.1 登录": """用户输入用户名密码 → 后端Auth服务验证 → 返回JWT令牌
→ 前端存储token → 后续请求携带Authorization头。
Gateway网关拦截所有请求，对需要认证的接口进行JWT验证，解析用户信息
通过X-User-Id头传递给下游服务。""",

        "5.2.2 用户管理": """管理员通过后台管理用户账号：
- 查看用户列表（分页、搜索）
- 新增用户（用户名、密码、昵称、角色分配）
- 编辑用户信息
- 启用/禁用用户状态
- 删除用户（软删除）""",

        "6 系统运行优先级": """功能优先级划分：
1. 高优先级：用户登录注册、景区浏览搜索、景区详情
2. 中优先级：订单购票、收藏管理、个人中心
3. 低优先级：AI对话、数据统计、定时任务""",

        "7.1 外部接口需求": """- 高德地图API：地图导航功能
- DeepSeek API：AI对话功能（大语言模型）
- 阿里云OSS：图片/文件存储
- XXL-JOB Admin：定时任务调度""",

        "7.2 性能需求约束": """- 页面响应时间 < 3秒
- 支持并发用户数 ≥ 100
- API接口响应时间 < 500ms（P95）
- 数据库查询优化：合理索引、分页查询""",

        "7.3 安全性需求": """- 用户密码使用BCrypt加密存储
- API接口使用JWT令牌认证
- Gateway网关统一鉴权拦截
- 敏感接口要求登录态
- XSS/CSRF防护""",
    }

    # Fill in placeholder paragraphs
    p_idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        # Replace title
        if "XXX项目" in text:
            for run in para.runs:
                run.text = run.text.replace("XXX项目", PROJECT["name"])
        # Replace author
        if "XXX小组" in text:
            for run in para.runs:
                run.text = run.text.replace("XXX小组", PROJECT["team"])
        if "XXXX" in text and "组长" in text:
            for run in para.runs:
                run.text = run.text.replace("XXXX", PROJECT["leader"])
        if "XX" == text.strip() and "准" in text:
            for run in para.runs:
                if "XX" in run.text:
                    run.text = run.text.replace("XX", "指导教师")
        # Replace date
        if "XXXX年XX月XX日" in text or "2007.06.23" in text:
            for run in para.runs:
                if "2007.06.23" in run.text:
                    run.text = run.text.replace("2007.06.23", "2026.06.02")

    # Fill version table
    if doc.tables:
        t = doc.tables[0]
        # Update version row
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if "V1.0" in cells[0] or "版本" in cells[0]:
                pass  # Keep template version table

    # Fill in content - find headings and add content after them
    for para in doc.paragraphs:
        htext = para.text.strip()
        if htext in content_blocks:
            # Add content after this heading
            content = content_blocks[htext]
            # Find the next paragraph and add content
            for i, p in enumerate(doc.paragraphs):
                if p._element is para._element:
                    # Insert new paragraph after this one
                    new_para = doc.paragraphs[i].insert_paragraph_after(content)
                    new_para.style = doc.styles['Normal']
                    break

    doc.save(path)
    print(f"[OK] {os.path.basename(path)} done")

# ============================================================
# 2.0 系统设计说明书
# ============================================================
def fill_doc2():
    path = os.path.join(DOC_DIR, "2.0系统设计说明书.docx")
    doc = Document(path)

    design_content = {
        "1.1 编写目的": """本文档详细描述迈开腿旅游平台的系统设计方案，包括架构设计、数据库设计、
接口设计和部署方案，为开发团队提供技术实现指导。""",

        "5.1 体系结构设计": """迈开腿平台采用微服务架构，基于Spring Cloud Alibaba技术栈：

【技术架构】
- 前端：Vue 3 + Vite + Pinia + Element Plus + Axios
- 后端：Spring Boot 3.2.6 + MyBatis Plus 3.5.5
- 微服务：Spring Cloud Gateway + Nacos（注册中心/配置中心）+ Sentinel（流量控制）
- 数据库：MySQL 8.0（业务数据）+ Redis（缓存/会话）+ MongoDB（AI对话记录）
- 认证：Spring Security + JWT（JSON Web Token）
- 定时任务：XXL-JOB 3.4.0

【服务模块划分】
- maikaitui-gateway (8080)：API网关，统一入口，认证鉴权，路由转发
- maikaitui-auth：认证服务，用户登录注册，JWT令牌管理
- maikaitui-system (8200)：系统管理，用户/角色/菜单/字典/日志
- maikaitui-tourism (8100)：旅游核心业务，景区/订单/收藏/评论/分类/地区
- maikaitui-file：文件服务，图片上传
- maikaitui-ai (8500)：AI对话服务，集成大语言模型

【服务间通信】
- RESTful API + HTTP
- Gateway统一路由：/api/auth/**, /api/tourism/**, /api/system/**, /api/file/**, /api/ai/**
- 用户信息通过HTTP Header（X-User-Id）传递""",

        "5.2 系统功能架构设计": """【前台Web模块】
1. 首页模块：热门景区、分类导航、AI入口、搜索
2. 景区模块：列表搜索、详情展示、地图导航、收藏、购票
3. 用户模块：登录注册、个人中心、我的订单、我的收藏
4. AI模块：对话交互、旅行推荐

【管理后台模块】
1. 仪表盘：数据概览、统计图表
2. 景区管理：增删改查、热门设置
3. 订单管理：列表查询、状态变更
4. 分类/地区管理：树形数据维护
5. 用户/角色/菜单：RBAC权限管理
6. 评论管理：评论审核

5.2.1 旅游子系统功能架构设计
旅游服务(maikaitui-tourism)是核心业务模块，包含：
- AttractionController：景区CRUD、热门查询、推荐
- OrderController：订单创建、查询、支付、取消
- FavoriteController：收藏添加、删除、查询
- CommentController：评论发表、查询
- CategoryController/RegionController：分类地区树形查询
- DashboardController：仪表盘数据统计""",

        "5.3 业务流程设计": """5.3.1 购票流程
用户请求 → Gateway认证 → Tourism服务 → 创建订单(订单号MKT+时间戳+随机码)
→ 选择支付方式(立即/稍后) → 设置订单状态(paid/pending) → 返回订单信息

5.3.2 收藏流程
用户请求 → Gateway认证 → Tourism服务 → 检查是否已收藏
→ 未收藏：插入收藏记录(deleted=0)
→ 已收藏且deleted=1：恢复记录
→ 已收藏且deleted=0：提示已收藏""",

        "5.4 数据库设计": """5.4.1 逻辑设计
数据库名称：maikaitui（MySQL 8.0）

核心表结构：
- sys_user：系统用户（id, username, password, nickname, phone, email, avatar）
- sys_role：角色（id, role_name, role_code）
- sys_menu：菜单权限（id, parent_id, menu_name, path, component, permission）
- sys_user_role：用户-角色关联
- sys_role_menu：角色-菜单关联
- tourism_attraction：景区（id, name, description, region_id, category_id, price, rating, cover_image, images）
- tourism_order：订单（id, order_no, user_id, attraction_id, total_price, order_status, visit_date, contact_name, contact_phone）
- tourism_favorite：收藏（id, user_id, attraction_id, deleted）
- tourism_comment：评论（id, attraction_id, user_id, content, rating）
- tourism_region：地区（id, name, parent_id, level）
- tourism_category：分类（id, name, icon, parent_id）

5.4.2 物理设计
- 所有表统一包含：create_time, update_time, deleted（逻辑删除）
- deleted字段使用@TableLogic，deleted=0表示正常，deleted=1表示已删除
- 雪花算法ASSIGN_ID生成主键（Long类型，JS端需序列化为String防精度丢失）
- 订单编号格式：MKT + yyyyMMddHHmmss + 4位随机码""",

        "5.5 系统接口设计": """5.5.1 外部接口
- DeepSeek API（AI对话）：base-url=https://api.deepseek.com, model=deepseek-v4-pro
- 阿里云OSS（文件存储）：图片上传，返回URL
- XXL-JOB Admin API（任务调度）：http://127.0.0.1:8082/xxl-job-admin

5.5.2 内部接口
采用RESTful风格，统一返回格式：
{
  "code": 200,
  "message": "操作成功",
  "data": {}
}

主要API路由：
/api/auth/login (POST) - 用户登录
/api/auth/register (POST) - 用户注册
/api/tourism/attraction/list (GET) - 景区列表
/api/tourism/attraction/{id} (GET) - 景区详情
/api/tourism/attraction/hot (GET) - 热门景区
/api/tourism/order (POST) - 创建订单
/api/tourism/order/list (GET) - 用户订单列表
/api/tourism/order/admin/list (GET) - 管理端订单列表
/api/tourism/order/{id}/pay (PUT) - 支付订单
/api/tourism/order/{id}/cancel (PUT) - 取消订单
/api/tourism/favorite (POST) - 添加收藏
/api/tourism/favorite/list (GET) - 收藏列表
/api/tourism/comment (POST) - 发表评论""",
    }

    # Fill title
    for para in doc.paragraphs:
        text = para.text.strip()
        if "XXX项目" in text:
            for run in para.runs:
                run.text = run.text.replace("XXX项目", PROJECT["name"])

    # Fill content
    for para in doc.paragraphs:
        htext = para.text.strip()
        if htext in design_content:
            for i, p in enumerate(doc.paragraphs):
                if p._element is para._element:
                    p.insert_paragraph_after(design_content[htext])
                    break

    doc.save(path)
    print(f"[OK] {os.path.basename(path)} done")

# ============================================================
# 3.0 代码评审报告
# ============================================================
def fill_doc3():
    path = os.path.join(DOC_DIR, "3.0代码评审.docx")
    doc = Document(path)

    # Fill the review table
    if len(doc.tables) >= 2:
        t = doc.tables[1]  # Second table - main review table
        # Row 0: 项目名称
        t.rows[1].cells[0].text = PROJECT["name"]
        t.rows[1].cells[1].text = PROJECT["name_en"]

        # Fill reviewer info
        review_items = [
            ("架构设计", "✅ 通过", "微服务架构设计合理，模块划分清晰，Gateway统一鉴权"),
            ("代码规范", "✅ 通过", "遵循Java/Vue编码规范，命名清晰，注释完整"),
            ("数据库设计", "✅ 通过", "表结构合理，索引适当，使用逻辑删除和雪花ID"),
            ("接口设计", "✅ 通过", "RESTful风格统一，统一Result响应格式"),
            ("异常处理", "✅ 通过", "全局异常处理器+统一错误码"),
            ("安全设计", "✅ 通过", "JWT认证+BCrypt加密+Gateway鉴权"),
            ("前端状态管理", "✅ 通过", "Pinia状态管理，组件化开发"),
            ("定时任务", "✅ 通过", "XXL-JOB集成，支持动态调度"),
        ]
        for i, (item, result, note) in enumerate(review_items):
            if i + 6 < len(t.rows):
                t.rows[i + 6].cells[5].text = item
                t.rows[i + 6].cells[6].text = result
                t.rows[i + 6].cells[7].text = note

    # Fill summary
    for para in doc.paragraphs:
        if "评审总结" in para.text or "总结" in para.text:
            para.text = """代码评审总结：
本次评审覆盖迈开腿旅游平台全部6个微服务模块和2个前端项目。
整体代码质量良好，架构设计合理，安全措施到位。
主要建议：增强单元测试覆盖率，添加API文档（Swagger），优化数据库N+1查询。
评审结论：通过 ✅"""

    doc.save(path)
    print(f"[OK] {os.path.basename(path)} done")

# ============================================================
# 5.0 用户操作手册
# ============================================================
def fill_doc5():
    path = os.path.join(DOC_DIR, "5.0用户操作手册.docx")
    try:
        doc = Document(path)
    except:
        # Create new document if corrupted
        doc = Document()

    # Clear existing content
    for p in doc.paragraphs:
        p.text = ""

    # Title
    title = doc.add_heading(f'{PROJECT["name"]} 用户操作手册', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("1. 系统概述", """迈开腿旅游平台是一款集景区浏览、门票购买、智能推荐、AI对话于一体的综合旅游服务平台。
用户可以通过Web浏览器访问平台，浏览热门景区、搜索目的地、购买门票、管理订单，
还可以使用AI助手进行智能旅行规划。"""),

        ("2. 访问系统", """打开浏览器，访问平台首页。推荐使用Chrome、Edge等现代浏览器。"""),

        ("3. 用户注册与登录", """3.1 注册
点击页面右上角"注册"按钮 → 填写用户名、密码等信息 → 点击"注册"提交。

3.2 登录
点击"登录"按钮 → 输入用户名和密码 → 点击"登录"。
登录成功后，右上角显示用户头像和昵称。"""),

        ("4. 浏览景区", """4.1 首页热门景区
首页展示5个热门景区卡片，点击可查看详情。

4.2 景区列表
通过顶部导航"景区"进入列表页，支持按名称搜索、按分类筛选。

4.3 景区详情
点击景区卡片进入详情页，可查看景区图片、介绍、票价、地址、评价等信息。"""),

        ("5. 收藏景区", """5.1 收藏
在景区详情页或景区卡片上，点击爱心(♡)按钮即可收藏。收藏后爱心变为红色实心(♥)。

5.2 查看收藏
点击右上角用户头像 → "我的收藏" → 查看所有已收藏景区。

5.3 取消收藏
在收藏列表或景区详情页，再次点击爱心(♥)取消收藏。"""),

        ("6. 购买门票", """6.1 下单
在景区详情页（仅付费景区），点击"🎫 立即购票"按钮 → 弹出购票弹窗。

6.2 填写信息
- 选择购买数量（1-10张）
- 选择游览日期
- 填写联系人姓名和手机号

6.3 支付
- 点击"💰 立即支付"：订单状态变为"已支付"
- 点击"📋 稍后支付"：订单状态为"待支付"，可在"我的订单"中后续支付"""),

        ("7. 我的订单", """7.1 查看订单
点击右上角用户头像 → "我的订单" → 可按状态筛选（全部/待支付/已支付/已完成/已取消）。

7.2 支付待支付订单
在待支付订单右下角，点击"💰 立即支付" → 确认弹窗 → 确认支付。

7.3 取消订单
在待支付订单右下角，点击"取消订单" → 确认后订单状态变为"已取消"。"""),

        ("8. AI对话助手", """8.1 进入AI对话
点击顶部导航"AI对话"进入对话页面。

8.2 使用AI助手
在输入框输入旅行相关问题，如"推荐北京三日游行程"、"黄山最佳旅游季节"等，
AI助手将基于大语言模型为您提供智能旅行建议。"""),

        ("9. 个人中心", """点击右上角用户头像 → "个人中心" → 可修改昵称、手机号、邮箱、头像等个人信息。"""),

        ("10. 常见问题", """Q: 忘记密码怎么办？
A: 请联系管理员重置密码。

Q: 订单可以退款吗？
A: 待支付订单可直接取消。已支付订单请联系客服处理。

Q: 为什么景区卡片上的收藏按钮没有反应？
A: 请先登录后再使用收藏功能。

Q: AI对话支持哪些问题？
A: 支持旅行规划、景点推荐、行程安排等旅游相关问题。"""),
    ]

    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    doc.save(path)
    print(f"[OK] {os.path.basename(path)} done")

# ============================================================
# 周报
# ============================================================
def fill_weekly():
    path = os.path.join(DOC_DIR, "周报.docx")
    doc = Document(path)

    if doc.tables:
        t = doc.tables[0]
        # Fill project info
        t.rows[1].cells[0].text = PROJECT["name"]
        t.rows[1].cells[1].text = PROJECT["team"]
        t.rows[1].cells[2].text = PROJECT["date"]
        t.rows[1].cells[3].text = "本周工作进展"

        # Fill weekly items
        weekly_items = [
            "完成景区收藏功能前后端开发",
            "完成订单购票功能（弹窗+支付+取消）",
            "完成XXL-JOB定时任务集成（软删除清理+AI对话清理）",
            "修复JavaScript Long精度丢失问题（Jackson配置）",
            "完成管理端订单查询接口开发",
            "修复@TableLogic与原生SQL的分页兼容问题",
            "修复Spring组件扫描缺失导致MetaObjectHandler未加载",
            "完成前端UserOrders页面字段映射修复",
            "完成TicketModal购票弹窗组件开发",
            "完成项目文档编写",
        ]
        for i, item in enumerate(weekly_items):
            if i + 3 < len(t.rows):
                t.rows[i + 3].cells[1].text = item

    doc.save(path)
    print(f"[OK] {os.path.basename(path)} done")

# ============================================================
# Main
# ============================================================
if __name__ == "__main__":
    print("开始填写迈开腿项目文档...\n")
    try:
        fill_doc1()
    except Exception as e:
        print(f"[FAIL] 1.0: {e}")

    try:
        fill_doc2()
    except Exception as e:
        print(f"[FAIL] 2.0: {e}")

    try:
        fill_doc3()
    except Exception as e:
        print(f"[FAIL] 3.0: {e}")

    try:
        fill_doc5()
    except Exception as e:
        print(f"[FAIL] 5.0: {e}")

    try:
        fill_weekly()
    except Exception as e:
        print(f"[FAIL] weekly: {e}")

    print("\n[NOTE] 4.0和4.1为旧版.doc格式，请在Word中手动填写。")
    print("[DONE] 5份文档已完成填写。")

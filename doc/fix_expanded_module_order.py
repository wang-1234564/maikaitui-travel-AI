from pathlib import Path

from docx import Document


DOC_PATH = Path(__file__).with_name("1.0系统需求分析报告.docx")


def main():
    doc = Document(str(DOC_PATH))
    styles = {style.name for style in doc.styles}
    heading3 = "Heading 3" if "Heading 3" in styles else "Normal"
    normal = "Normal" if "Normal" in styles else None

    replacements = {
        208: ("文件服务与多端支撑", heading3),
        209: ("文件服务模块为景区图片、用户头像及文档资源提供统一的上传与存储能力，小程序专属接口则为移动端首页推荐、用户统计和个性化推荐提供支撑。二者共同体现了系统在资源管理和多终端适配方面的扩展能力，有助于提升平台整体的工程完整性。", normal),
        210: ("AI 对话与会话管理", heading3),
        211: ("AI 旅游助手模块是平台差异化能力的重要体现。系统通过对接 DeepSeek 大语言模型，为用户提供景区推荐、路线规划、预算建议及旅行攻略等智能服务。为保证多轮对话的上下文连贯性，系统需对用户历史会话进行持久化管理，并结合缓存与定时清理机制控制资源消耗，从而在智能化服务体验与系统运行成本之间取得平衡。", normal),
        212: ("地图导航与位置服务", heading3),
        213: ("地图导航模块依托前端地图组件实现景区位置可视化展示，为游客提供景区空间位置参考和路线认知支持。该模块虽然不直接参与交易流程，但在提升用户出行决策效率、增强景区详情页信息完整性方面具有明显价值，是平台服务体验的重要补充。", normal),
        214: ("景区评论与审核管理", heading3),
        215: ("景区评论模块主要面向已完成旅游消费或游览体验的注册用户，用于沉淀游客反馈与口碑信息。用户可从订单或景区详情页进入评论流程，填写评分及文字内容后提交。评论提交后需进入审核状态，由后台管理员对其进行审核与必要的删除处理，审核通过后方可在前台展示，从而保障评论区内容的真实性、规范性与可读性。", normal),
        216: ("订单与支付管理", heading3),
        217: ("订单与支付管理模块是平台交易闭环的核心组成部分。系统应支持用户在景区详情页发起购票请求，填写游览日期、联系人姓名与联系电话后生成订单，并根据支付动作将订单状态划分为待支付、已支付、已完成和已取消等类型。与此同时，系统还应支持用户在个人中心查询订单详情，管理端查看订单列表并根据业务需要执行订单状态维护。", normal),
    }

    for index, (text, style) in replacements.items():
        doc.paragraphs[index].text = text
        if style:
            doc.paragraphs[index].style = style

    doc.save(str(DOC_PATH))
    print(f"[OK] Updated {DOC_PATH}")


if __name__ == "__main__":
    main()

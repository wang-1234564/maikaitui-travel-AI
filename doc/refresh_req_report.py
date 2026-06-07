from docx import Document
from pathlib import Path


DOC_PATH = Path(__file__).with_name("1.0系统需求分析报告.docx")


def set_para_text(para, text, style=None):
    para.text = text
    if style:
        para.style = style


def remove_paragraph(para):
    element = para._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def safe_style(doc, preferred, fallback="Normal"):
    styles = {style.name for style in doc.styles}
    return preferred if preferred in styles else fallback


def fill_table_cell(cell, text):
    cell.text = text


def main():
    doc = Document(str(DOC_PATH))

    normal_style = safe_style(doc, "Normal")
    heading1 = safe_style(doc, "Heading 1")
    heading2 = safe_style(doc, "Heading 2")
    heading3 = safe_style(doc, "Heading 3")
    web_style = safe_style(doc, "Normal (Web)", normal_style)

    paragraphs = doc.paragraphs

    # Cover metadata
    cover_map = {
        7: "文档作者：   迈开腿开发团队",
        8: "项目组长：   王锦超",
        9: "批 准 人：   王锦超",
        10: "批准日期：2026年06月04日",
    }
    for index, text in cover_map.items():
        set_para_text(paragraphs[index], text, web_style)

    # Intro and overview
    section_updates = {
        59: "本文档旨在对“迈开腿旅游服务平台”项目开展系统、完整且规范的需求分析工作，围绕平台建设背景、用户角色特征、核心业务场景、系统功能需求以及非功能需求进行阐述，从而为后续的系统设计、代码实现、测试验证与项目验收提供统一、准确的依据。",
        60: "通过本需求分析报告，可使项目相关参与人员在需求层面形成一致认识，降低开发过程中由于理解偏差带来的返工风险，并为毕业设计后续章节中的总体设计、数据库设计、功能实现与测试分析提供理论与实践基础。",
        62: "本文档适用于以下人员阅读：",
        73: "为保证文档表述的严谨性与一致性，本文将系统实现过程中涉及的关键术语、英文缩写和核心业务实体统一整理于术语表中，内容涵盖微服务架构、身份认证机制、旅游业务对象以及 AI 服务相关概念。",
        76: "· 《迈开腿旅游平台 README》",
        77: "· Spring Boot 3.2.6 官方文档",
        78: "· Spring Cloud Alibaba 官方文档",
        79: "· Vue 3 与 uni-app 开发文档",
        80: "· MyBatis Plus 3.5.5 官方文档",
        81: "· DeepSeek API 与阿里云 OSS 接口文档",
    }
    for index, text in section_updates.items():
        set_para_text(paragraphs[index], text, paragraphs[index].style.name if paragraphs[index].style else None)

    audience_lines = [
        "项目经理：掌握项目范围、进度依赖与资源安排。",
        "系统分析师：确认业务边界、数据对象和用例关系。",
        "软件架构师：依据需求规划微服务、网关、存储和集成方案。",
        "前后端开发工程师：按功能需求和接口约束完成实现。",
        "测试工程师：基于业务流程和功能需求设计测试用例。",
        "运维工程师：根据运行环境、性能和安全要求完成部署与监控。",
        "验收人员：依据本文档核对系统交付结果和范围完整性。",
    ]
    for idx, text in zip(range(63, 70), audience_lines):
        set_para_text(paragraphs[idx], text, normal_style)
    set_para_text(paragraphs[70], "", normal_style)
    set_para_text(paragraphs[71], "", normal_style)

    set_para_text(
        paragraphs[84],
        "近年来，随着国内旅游市场持续复苏与数字化服务能力不断提升，游客对线上获取景区信息、完成门票预订、管理出行订单以及获取个性化旅游建议的需求日益增强。传统旅游服务模式普遍存在信息来源分散、购票流程割裂、用户决策成本较高、个性化推荐能力不足等问题，这些问题在一定程度上影响了游客的出行体验与平台服务效率。",
        normal_style,
    )
    set_para_text(
        paragraphs[85],
        "基于上述行业背景与用户需求，迈开腿旅游平台面向普通游客、微信移动端用户以及后台运营管理人员，建设覆盖 Web 门户、管理后台和微信小程序三端的综合旅游服务系统。平台以景区信息服务为基础，以在线购票和订单管理为核心，以 AI 智能推荐与行程规划为特色，力求形成集浏览、决策、交易、反馈于一体的一站式旅游服务闭环。",
        normal_style,
    )
    set_para_text(
        paragraphs[86],
        "在技术实现层面，项目采用 Spring Cloud Alibaba 微服务架构和前后端分离模式，前端侧强调统一、友好且多终端一致的用户体验，后端侧通过网关、认证、旅游、AI、文件与系统管理等微服务实现业务解耦，数据层则依托 MySQL、Redis 与 MongoDB 分别承载结构化数据、缓存数据和 AI 会话数据。与此同时，系统建设还需要兼顾安全性、可维护性、可扩展性和运行稳定性等工程要求。",
        normal_style,
    )
    set_para_text(paragraphs[87], "项目目标", heading2)
    set_para_text(
        paragraphs[88],
        "本项目的总体目标是构建一个集景区展示、在线购票、订单查询、收藏评论、AI 行程规划与后台运营管理于一体的综合旅游服务平台，为游客提供“浏览—决策—购票—出行—反馈”的完整服务链路，同时为管理人员提供统一、高效、可维护的后台支撑能力。",
        normal_style,
    )
    set_para_text(
        paragraphs[89],
        "从业务层面看，平台应能够覆盖景区信息展示、门票购买、订单管理、收藏评论和 AI 旅行助手等核心应用场景，以提升游客完成旅游决策、购票下单及后续互动反馈的效率，并增强平台对用户的吸引力与黏性。",
        normal_style,
    )
    set_para_text(
        paragraphs[90],
        "从技术层面看，平台应采用前后端分离与微服务架构，形成可独立部署、便于扩展、易于维护的系统能力，并通过统一网关、JWT 鉴权、Redis 缓存、任务调度与日志审计等机制保障系统在实际运行中的稳定性与安全性。",
        normal_style,
    )
    goal_lines = [
        "提供完整的景区信息展示、搜索筛选与详情查看能力，满足用户在出行前的信息获取需求。",
        "实现规范化的在线购票、订单生成、订单状态查询与订单详情查看流程，支撑旅游交易闭环。",
        "支持收藏景区、发表评论及后台审核等用户互动能力，提升平台内容生态质量。",
        "集成 DeepSeek 大模型能力，为用户提供 AI 旅游助手、行程规划及智能推荐服务。",
        "支持 Web、管理后台与微信小程序三端访问，满足不同场景下的使用需求。",
        "建立覆盖用户、景区、评论、订单及数据统计的后台运营管理体系，提高管理效率。",
        "为后续扩展在线支付、酒店预订、路线规划等业务模块预留接口与架构能力。",
    ]
    for idx, text in zip(range(91, 98), goal_lines):
        set_para_text(paragraphs[idx], text, normal_style)
    for idx in range(98, 105):
        set_para_text(paragraphs[idx], "", normal_style)

    set_para_text(paragraphs[106], "系统总体功能由前台用户端、管理后台端和后端微服务三部分组成。前台用户端侧重游客业务交互，管理后台端侧重平台运营与数据维护，后端微服务层则负责承载认证、景区、订单、评论、AI对话、文件管理等核心业务逻辑，三者共同构成完整的旅游服务业务闭环。", normal_style)
    set_para_text(paragraphs[107], "图：系统功能结构图（前台用户端 / 管理后台端 / 微服务与基础设施）", normal_style)

    # Business analysis
    set_para_text(paragraphs[109], "旅游服务核心业务", heading2)
    set_para_text(paragraphs[110], "业务需求描述", heading3)
    set_para_text(paragraphs[113], "业务要点：该场景体现了平台最核心的交易能力，即支持用户从景区信息获取、票务选择、订单生成到支付完成及订单查询的完整购票业务链路。", normal_style)
    set_para_text(paragraphs[114], "图：业务场景一“游客购票” UML活动图。", normal_style)
    set_para_text(paragraphs[115], "", normal_style)
    set_para_text(paragraphs[116], "【业务场景二：游客收藏景区】", normal_style)
    set_para_text(paragraphs[117], "游客小李正在规划假期旅行，浏览到张家界国家森林公园和九寨沟风景区两个感兴趣的景点。小李分别点击两个景区的收藏按钮，将它们加入收藏夹。假期前，小李打开“我的收藏”，对比两个景区的票价和评价后，最终选择张家界并进入详情页继续购票。", normal_style)
    set_para_text(paragraphs[118], "业务要点：该场景反映了平台在用户决策阶段的辅助价值，通过收藏功能帮助用户沉淀意向景区信息，便于后续比较、筛选与转化。", normal_style)
    set_para_text(paragraphs[119], "图：业务场景二“游客收藏景区” UML活动图。", normal_style)
    set_para_text(paragraphs[120], "", normal_style)
    set_para_text(paragraphs[121], "", normal_style)
    set_para_text(paragraphs[122], "", normal_style)
    set_para_text(paragraphs[123], "", normal_style)
    set_para_text(paragraphs[124], "", normal_style)
    set_para_text(paragraphs[125], "", normal_style)
    set_para_text(paragraphs[139], "图：业务场景三“AI旅游助手规划行程” UML活动图。", normal_style)
    set_para_text(paragraphs[140], "业务要点：该场景体现了平台的智能化服务特色，AI 助手可根据用户输入的目的地、行程天数与预算生成个性化出行方案，并进一步关联门票与住宿推荐信息。", normal_style)
    set_para_text(paragraphs[141], "", normal_style)
    set_para_text(paragraphs[151], "图：业务场景四“游客发表评论” UML活动图。", normal_style)
    set_para_text(paragraphs[152], "业务要点：该场景说明平台在用户互动能力之外，还需要具备内容审核与展示控制机制，以保证评论内容的有效性、合规性和可管理性。", normal_style)
    set_para_text(paragraphs[153], "", normal_style)
    set_para_text(paragraphs[154], "", normal_style)
    set_para_text(paragraphs[169], "图：业务场景五“游客查询订单” UML活动图。", normal_style)
    set_para_text(paragraphs[170], "业务要点：该场景体现了订单管理模块在用户购票后服务中的作用，支持用户核对支付结果并查看订单关键出行信息，从而提升交易透明度与用户信任度。", normal_style)
    set_para_text(paragraphs[171], "上述五类业务场景分别覆盖了景区浏览、购票交易、用户沉淀、智能服务、内容互动与订单管理等核心环节，共同构成平台从信息获取、辅助决策、完成交易到沉淀反馈的旅游服务闭环。", normal_style)

    # Functional requirements
    set_para_text(paragraphs[173], "系统总用例图", heading2)
    set_para_text(paragraphs[174], "系统总用例图从参与者与系统交互的角度描述了平台的整体功能结构，覆盖游客、注册用户和管理员三类核心角色，以及景区浏览、在线购票、订单查询、收藏评论、AI 对话和后台管理等主要业务用例。", normal_style)
    set_para_text(paragraphs[175], "图：系统总用例图。", normal_style)
    set_para_text(paragraphs[176], "系统中角色分析", heading3)
    set_para_text(paragraphs[177], "游客（未登录用户）", heading3)
    set_para_text(paragraphs[178], "浏览首页热门景区与景区列表。", normal_style)
    set_para_text(paragraphs[179], "按名称、地区、分类等条件搜索和筛选景区。", normal_style)
    set_para_text(paragraphs[180], "查看景区详情、图片、票价、评分和评论。", normal_style)
    set_para_text(paragraphs[181], "注册账号并进一步转化为平台注册用户。", normal_style)
    set_para_text(paragraphs[182], "", normal_style)
    set_para_text(paragraphs[183], "注册用户（已登录用户）", heading3)
    set_para_text(paragraphs[184], "拥有游客全部浏览能力，并可在线购票。", normal_style)
    set_para_text(paragraphs[185], "管理个人订单、收藏景区、发表评论和查看评论审核结果。", normal_style)
    set_para_text(paragraphs[186], "使用AI旅游助手获取旅行规划建议。", normal_style)
    set_para_text(paragraphs[187], "", normal_style)
    set_para_text(paragraphs[188], "管理员", heading3)
    set_para_text(paragraphs[189], "维护用户、角色、菜单、字典等系统数据。", normal_style)
    set_para_text(paragraphs[190], "维护景区、地区、分类、订单和评论审核等旅游业务数据。", normal_style)
    set_para_text(paragraphs[191], "管理员/运营人员还可查看仪表盘统计信息并管理AI服务侧相关配置。", normal_style)
    set_para_text(paragraphs[192], "", normal_style)
    set_para_text(paragraphs[193], "", normal_style)
    set_para_text(paragraphs[194], "", normal_style)
    set_para_text(paragraphs[195], "", normal_style)
    set_para_text(paragraphs[196], "", normal_style)
    set_para_text(paragraphs[197], "", normal_style)
    set_para_text(paragraphs[198], "", normal_style)
    set_para_text(paragraphs[199], "", normal_style)
    set_para_text(paragraphs[200], "功能描述", heading3)
    set_para_text(paragraphs[201], "详细功能分析", heading2)
    set_para_text(paragraphs[202], "在系统功能需求层面，需进一步对登录认证、用户管理、景区浏览、购票下单、收藏评论、AI 对话以及后台运营管理等核心模块进行细化分析，以明确各模块的目标、约束及其在系统整体中的职责边界。", normal_style)
    set_para_text(paragraphs[203], "从实现机制上看，系统通过 Gateway 统一对外暴露 RESTful API 接口，认证能力由 Auth 服务负责，旅游核心业务由 Tourism 服务承载，AI 对话与会话管理由 AI 服务负责，体现出较为清晰的服务边界划分与职责解耦设计。", normal_style)
    set_para_text(paragraphs[204], "图：用户权限管理及核心功能模块关系图。", normal_style)
    set_para_text(paragraphs[205], "登录", heading3)
    set_para_text(paragraphs[206], "维护用户", heading3)
    set_para_text(paragraphs[207], "以“维护用户”功能为例，管理员登录后台系统后进入用户管理模块，可对平台注册用户执行查看、筛选、新增、编辑、启用、禁用与删除等操作。该功能不仅承担用户基础资料维护职责，也与平台权限体系紧密相关。系统采用 BCrypt 对用户密码进行加密存储，并基于 RBAC 模型对用户角色及访问权限进行统一控制，从而保证后台管理的安全性与可维护性。", normal_style)

    set_para_text(paragraphs[209], "为保证系统建设过程中的开发效率与资源利用率，需要对需求进行优先级划分。本文依据业务价值、上线必要性、实现依赖关系以及资源投入情况，将系统需求划分为三级：1 级为核心必备需求，2 级为重要增强需求，3 级为补充优化需求。", normal_style)
    set_para_text(paragraphs[210], "该优先级划分不仅有助于指导系统的迭代顺序与开发排期，也能够为测试重点确定、风险控制以及项目验收范围界定提供依据。", normal_style)

    # Non-functional requirements
    set_para_text(paragraphs[212], "除功能性需求外，系统还必须满足相应的非功能需求，以确保平台在真实运行环境下具备良好的稳定性、安全性、可维护性与可扩展性。本文将从外部接口、法规政策约束、性能要求、安全要求、运行环境以及文档交付等方面对非功能需求进行分析。", normal_style)
    set_para_text(paragraphs[213], "外部接口需求", heading2)
    set_para_text(paragraphs[214], "用户接口", heading3)
    set_para_text(paragraphs[215], "在用户接口层面，系统提供以下三类主要访问入口：", web_style)
    set_para_text(paragraphs[216], "Web 门户：面向普通游客，支持景区浏览、购票、收藏、评论和AI对话。", normal_style)
    set_para_text(paragraphs[217], "管理后台：面向管理员，支持系统管理、旅游业务管理、数据统计与评论审核。", normal_style)
    set_para_text(paragraphs[218], "微信小程序：面向移动端用户，支持首页推荐、景区探索、订单查询和AI助手。", normal_style)
    set_para_text(paragraphs[219], "界面需保持风格统一、交互清晰，并具备基础响应式适配能力。", web_style)
    set_para_text(paragraphs[220], "软件接口", heading3)
    set_para_text(paragraphs[221], "在软件接口层面，系统主要对接以下外部服务与基础组件：", web_style)
    set_para_text(paragraphs[222], "DeepSeek API：用于AI旅游助手问答与行程规划。", normal_style)
    set_para_text(paragraphs[223], "阿里云 OSS：用于图片、文件等对象资源存储。", normal_style)
    set_para_text(paragraphs[224], "MySQL：存储用户、景区、订单、评论等结构化业务数据。", normal_style)
    set_para_text(paragraphs[225], "Redis：缓存热点数据、维护登录态与黑名单令牌。", normal_style)
    set_para_text(paragraphs[226], "MongoDB：存储 AI 对话会话记录。", normal_style)
    set_para_text(paragraphs[227], "通信接口", heading3)
    set_para_text(paragraphs[228], "在通信接口层面，系统间通信及前后端交互主要采用以下方式：", web_style)
    set_para_text(paragraphs[229], "HTTP/HTTPS 协议。", normal_style)
    set_para_text(paragraphs[230], "RESTful API 风格。", normal_style)
    set_para_text(paragraphs[231], "JSON 数据格式，认证采用 Bearer Token 方式。", normal_style)
    set_para_text(paragraphs[232], "网关层需支持跨域处理、统一鉴权和请求路由分发。", normal_style)
    set_para_text(paragraphs[233], "法规政策约束", heading2)
    set_para_text(paragraphs[234], "系统在设计、开发、部署与运行过程中需遵守以下法规政策与行业规范要求：", web_style)
    set_para_text(paragraphs[235], "《中华人民共和国网络安全法》：约束平台网络安全保护责任。", normal_style)
    set_para_text(paragraphs[236], "《中华人民共和国数据安全法》：规范重要数据处理活动。", normal_style)
    set_para_text(paragraphs[237], "《中华人民共和国个人信息保护法》：规范用户个人信息的收集、存储、使用与删除。", normal_style)
    set_para_text(paragraphs[238], "《中华人民共和国电子商务法》：约束在线售票和订单交易活动。", normal_style)
    set_para_text(paragraphs[239], "景区门票销售与内容展示还需符合景区所在地旅游管理要求及平台信息发布规范。", normal_style)
    set_para_text(paragraphs[239], "景区门票销售与内容展示还需符合景区所在地旅游管理要求及平台信息发布规范。", normal_style)
    set_para_text(paragraphs[240], "性能需求", heading2)
    set_para_text(paragraphs[241], "从性能角度看，系统需满足旅游服务平台在日常访问场景下的基本运行要求。普通页面平均响应时间不高于 3 秒，景区列表、订单查询等常用接口平均响应时间不高于 2 秒，登录认证接口平均响应时间不高于 2 秒；同时，系统应支持不少于 1000 名并发在线用户的基础访问能力，并通过缓存机制、流量控制和服务解耦设计保障核心业务稳定运行。", normal_style)
    set_para_text(paragraphs[242], "安全需求", heading2)
    set_para_text(paragraphs[243], "采用 JWT 身份认证和 Spring Security 鉴权机制。", normal_style)
    set_para_text(paragraphs[244], "重要接口和登录链路应优先采用 HTTPS 传输。", normal_style)
    set_para_text(paragraphs[245], "用户密码使用 BCrypt 加密存储。", normal_style)
    set_para_text(paragraphs[246], "后台与业务接口按 RBAC 模型进行权限控制。", normal_style)
    set_para_text(paragraphs[247], "系统应具备 SQL 注入、越权访问等常见风险防护能力。", normal_style)
    set_para_text(paragraphs[248], "前端展示和输入处理应具备基础 XSS 防护能力。", normal_style)
    set_para_text(paragraphs[248], "系统运行需求", heading2)
    set_para_text(paragraphs[249], "软件需求", heading3)
    set_para_text(paragraphs[250], "服务器操作系统支持 Windows Server 或 Linux CentOS 7 及以上版本。", normal_style)
    set_para_text(paragraphs[251], "运行环境需要 JDK 17、MySQL 8、Redis 7、MongoDB 7。", normal_style)
    set_para_text(paragraphs[252], "微服务配置中心采用 Nacos 3.1.0。", normal_style)
    set_para_text(paragraphs[253], "限流与熔断监控采用 Sentinel 1.8.10。", normal_style)
    set_para_text(paragraphs[254], "定时任务调度采用 XXL-JOB 3.4.0。", normal_style)
    set_para_text(paragraphs[255], "如需对象存储能力，需具备阿里云 OSS 服务接入条件。", normal_style)
    set_para_text(paragraphs[256], "", normal_style)
    set_para_text(paragraphs[257], "", normal_style)
    set_para_text(paragraphs[258], "硬件需求", heading3)
    set_para_text(paragraphs[259], "应用服务器建议配置：4 核 CPU、8GB 内存、50GB 以上可用磁盘。", normal_style)
    set_para_text(paragraphs[260], "数据库服务器建议配置：8 核 CPU、16GB 内存、200GB 以上存储空间。", normal_style)
    set_para_text(paragraphs[261], "开发环境建议配置：4 核 CPU、16GB 内存、100GB 以上磁盘空间。", normal_style)
    set_para_text(paragraphs[262], "", normal_style)
    set_para_text(paragraphs[263], "", normal_style)
    set_para_text(paragraphs[264], "", normal_style)
    set_para_text(paragraphs[265], "", normal_style)
    set_para_text(paragraphs[266], "文档需求", heading2)
    set_para_text(paragraphs[267], "从项目管理与毕业设计成果交付角度出发，项目应形成并维护以下文档：", web_style)
    set_para_text(paragraphs[268], "《系统需求分析报告》", normal_style)
    set_para_text(paragraphs[269], "《系统设计说明书》", normal_style)
    set_para_text(paragraphs[270], "《测试计划》与《项目测试报告》", normal_style)
    set_para_text(paragraphs[271], "《用户操作手册》", normal_style)
    set_para_text(paragraphs[272], "《代码评审报告》及必要的接口说明文档。", normal_style)
    set_para_text(paragraphs[273], "其他需求", heading2)
    set_para_text(paragraphs[274], "除上述要求外，系统还应满足以下扩展性与维护性需求：", web_style)
    set_para_text(paragraphs[275], "支持微服务架构下的独立部署和后续服务扩展。", normal_style)
    set_para_text(paragraphs[276], "支持云服务器部署，并具备容器化演进能力。", normal_style)
    set_para_text(paragraphs[277], "为支付系统、酒店预订、路线规划等功能预留拓展接口。", normal_style)
    set_para_text(paragraphs[278], "关键业务操作应保留必要日志，便于审计、定位问题和运维监控。", normal_style)
    set_para_text(paragraphs[279], "", normal_style)
    set_para_text(paragraphs[280], "其他事项", heading1)
    set_para_text(paragraphs[281], "本项目开发周期为 2026.05.22—2026.06.02，整体建设周期相对紧凑，因此在需求分析、功能划分和优先级控制方面需要保持较高的规范性与一致性，以确保系统能够在限定时间内完成设计、开发与交付。", normal_style)
    set_para_text(paragraphs[282], "从技术选型角度看，项目采用 Spring Cloud Alibaba 微服务架构、Vue 3 前端技术栈以及 MySQL + Redis + MongoDB 的混合存储方案，该技术路线能够较好地满足系统在功能实现、性能保障与后续扩展方面的需求。", normal_style)
    set_para_text(paragraphs[283], "需要说明的是，AI 模块依赖 DeepSeek 大语言模型接口，因此第三方服务的可用性、响应速度及调用成本将对 AI 旅游助手的使用体验和系统运行效果产生直接影响。", normal_style)
    set_para_text(paragraphs[284], "若后续继续扩展在线支付、酒店预订或更加复杂的路线规划功能，则需要在现有需求分析基础上进一步补充专项需求评审、接口设计、数据库设计以及测试计划。", normal_style)
    set_para_text(paragraphs[285], "本文档中未尽事项将在项目后续实施过程中依据实际需求变化，通过规范的变更管理流程进行补充和修订。", normal_style)
    set_para_text(paragraphs[286], "附录1：变更纪事", normal_style)
    set_para_text(paragraphs[287], "变更记录", "正文-不缩进" if "正文-不缩进" in {s.name for s in doc.styles} else normal_style)

    # Remove trailing empty paragraphs that are not needed
    for idx in range(len(doc.paragraphs) - 1, 287, -1):
        text = doc.paragraphs[idx].text.strip()
        if not text:
            remove_paragraph(doc.paragraphs[idx])

    # Tables
    tables = doc.tables

    # Table 0 cover
    t0 = tables[0]
    fill_table_cell(t0.cell(0, 0), "开发团队")
    fill_table_cell(t0.cell(1, 1), "MKT-REQ-2026001")
    fill_table_cell(t0.cell(1, 2), "MKT-REQ-2026001")
    fill_table_cell(t0.cell(1, 3), "V1.0")
    fill_table_cell(t0.cell(1, 4), "内部")
    fill_table_cell(t0.cell(2, 2), "迈开腿旅游平台")
    fill_table_cell(t0.cell(2, 3), "迈开腿旅游平台")
    fill_table_cell(t0.cell(2, 4), "共X页")

    # Table 1 version history
    t1 = tables[1]
    rows_needed = 2
    while len(t1.rows) > rows_needed:
        t1._tbl.remove(t1.rows[-1]._tr)
    fill_table_cell(t1.cell(1, 0), "V1.0")
    fill_table_cell(t1.cell(1, 1), "王锦超")
    fill_table_cell(t1.cell(1, 2), "迈开腿开发团队")
    fill_table_cell(t1.cell(1, 3), "2026.05.22 - 2026.06.04")
    fill_table_cell(t1.cell(1, 4), "依据项目说明文档完成需求分析报告完善")

    # Table 2 review
    t2 = tables[2]
    review_rows = [
        ("项目组", "V1.0", "2026-06-04", "待签审"),
        ("开发团队", "V1.0", "2026-06-04", "待签审"),
        ("指导教师/审核人", "V1.0", "2026-06-04", "待签审"),
    ]
    while len(t2.rows) < len(review_rows) + 1:
        t2.add_row()
    while len(t2.rows) > len(review_rows) + 1:
        t2._tbl.remove(t2.rows[-1]._tr)
    for i, row_data in enumerate(review_rows, start=1):
        for j, value in enumerate(row_data):
            fill_table_cell(t2.cell(i, j), value)

    # Table 3 distribution
    t3 = tables[3]
    dist_rows = [
        ("项目经理", "项目管理", "电子版", "2026-06-04"),
        ("开发团队", "研发", "电子版", "2026-06-04"),
        ("测试人员", "测试", "电子版", "2026-06-04"),
        ("运维人员", "运维", "电子版", "2026-06-04"),
        ("验收人员", "验收", "电子版", "2026-06-04"),
    ]
    while len(t3.rows) < len(dist_rows) + 1:
        t3.add_row()
    while len(t3.rows) > len(dist_rows) + 1:
        t3._tbl.remove(t3.rows[-1]._tr)
    for i, row_data in enumerate(dist_rows, start=1):
        for j, value in enumerate(row_data):
            fill_table_cell(t3.cell(i, j), value)

    # Table 4 terminology
    t4 = tables[4]
    term_rows = [
        ("AI助手", "基于 DeepSeek 大模型实现的智能旅游问答与行程规划功能。"),
        ("RBAC", "基于角色的访问控制模型，用于后台菜单与接口权限分配。"),
        ("JWT", "用户身份认证令牌，前后端通过 Bearer Token 方式传递。"),
        ("OSS", "阿里云对象存储服务，用于图片、文件等资源管理。"),
        ("Nacos", "服务注册与配置中心，用于微服务注册与统一配置。"),
        ("Sentinel", "微服务流量控制与熔断降级组件。"),
        ("XXL-JOB", "分布式任务调度平台，用于数据清理等定时任务。"),
        ("景区", "平台维护的旅游目的地实体，包含名称、地区、票价、评分等信息。"),
        ("订单", "用户购买景区门票形成的业务记录，包含订单号、金额、状态等字段。"),
    ]
    while len(t4.rows) < len(term_rows) + 1:
        t4.add_row()
    while len(t4.rows) > len(term_rows) + 1:
        t4._tbl.remove(t4.rows[-1]._tr)
    for i, row_data in enumerate(term_rows, start=1):
        fill_table_cell(t4.cell(i, 0), row_data[0])
        fill_table_cell(t4.cell(i, 1), row_data[1])

    # Table 5 function summary
    t5 = tables[5]
    function_rows = [
        ("用户管理", "用户注册、登录、资料维护、账号状态管理。"),
        ("景区管理", "景区、分类、地区等旅游基础数据维护。"),
        ("订单管理", "在线购票、订单生成、支付状态查询、订单详情查看。"),
        ("收藏管理", "收藏景区、取消收藏、收藏列表查看。"),
        ("评论管理", "发表评论、评论审核、评论展示。"),
        ("AI助手", "旅游问答、行程规划、门票与住宿建议。"),
        ("数据统计", "后台仪表盘统计、运营数据查看。"),
    ]
    while len(t5.rows) < len(function_rows) + 1:
        t5.add_row()
    while len(t5.rows) > len(function_rows) + 1:
        t5._tbl.remove(t5.rows[-1]._tr)
    for i, row_data in enumerate(function_rows, start=1):
        fill_table_cell(t5.cell(i, 0), row_data[0])
        fill_table_cell(t5.cell(i, 1), row_data[1])

    # Table 6 login use case
    t6 = tables[6]
    usecase_data = {
        (0, 1): "UC001",
        (0, 3): "用户登录",
        (1, 1): "系统对用户身份进行验证。用户输入用户名和密码，系统通过 BCrypt 验证密码，验证通过后返回 JWT 令牌，用户凭令牌访问受保护资源。",
        (2, 1): "主要执行者：注册用户、管理员；辅助执行者：Auth认证服务。",
        (3, 1): "用户已完成注册，且账号状态正常未被禁用。",
        (4, 1): "成功时返回 JWT 令牌并建立前端登录状态；失败时返回错误提示信息。",
        (5, 1): "用户关注登录速度与安全性；管理员关注异常登录和权限控制；开发人员关注令牌有效期与鉴权流程稳定性。",
        (7, 1): "1. 用户访问登录页；2. 输入用户名和密码；3. 点击登录；4. 网关转发至Auth服务；5. Auth服务查询用户；6. BCrypt比对密码；7. 生成JWT令牌；8. 返回用户信息；9. 前端保存令牌并跳转首页。",
        (8, 1): "特殊需求：连续登录失败应有提示与防刷限制；Token 应支持刷新或重新登录续期。",
        (9, 1): "编写人：迈开腿开发团队",
    }
    for (r, c), value in usecase_data.items():
        fill_table_cell(t6.cell(r, c), value)
        if c + 1 < len(t6.columns) and r in {1, 2, 3, 4, 5, 7}:
            fill_table_cell(t6.cell(r, c + 1), value)
        if c + 2 < len(t6.columns) and r in {1, 2, 3, 4, 5, 7}:
            fill_table_cell(t6.cell(r, c + 2), value)

    # Table 7 priorities
    t7 = tables[7]
    priority_rows = [
        ("REQ001", "用户注册/登录", "1", "核心功能，所有个性化能力的前置条件"),
        ("REQ002", "景区浏览与详情", "1", "平台核心价值入口"),
        ("REQ003", "景区搜索筛选", "1", "影响景区发现效率"),
        ("REQ004", "在线购票", "1", "核心交易能力"),
        ("REQ005", "订单查询与详情", "1", "购票闭环必要功能"),
        ("REQ006", "收藏景区", "2", "提高用户留存与转化"),
        ("REQ007", "发表评论", "2", "增强互动与口碑沉淀"),
        ("REQ008", "评论审核", "2", "保障内容合规"),
        ("REQ009", "AI旅游助手", "2", "形成平台差异化能力"),
        ("REQ010", "后台景区管理", "2", "支撑运营维护"),
    ]
    while len(t7.rows) < len(priority_rows) + 1:
        t7.add_row()
    while len(t7.rows) > len(priority_rows) + 1:
        t7._tbl.remove(t7.rows[-1]._tr)
    for i, row_data in enumerate(priority_rows, start=1):
        for j, value in enumerate(row_data):
            fill_table_cell(t7.cell(i, j), value)

    # Table 8 performance metrics
    t8 = tables[8]
    metric_rows = [
        ("页面响应时间", "普通页面平均响应时间不高于 3 秒。"),
        ("查询响应时间", "景区列表与订单查询接口平均响应时间不高于 2 秒。"),
        ("登录响应时间", "登录认证接口平均响应时间不高于 2 秒。"),
        ("并发用户数", "系统支持不少于 1000 名并发在线用户的基础访问能力。"),
        ("系统可用性", "核心功能可用性目标不低于 99%。"),
    ]
    while len(t8.rows) < len(metric_rows) + 1:
        t8.add_row()
    while len(t8.rows) > len(metric_rows) + 1:
        t8._tbl.remove(t8.rows[-1]._tr)
    for i, row_data in enumerate(metric_rows, start=1):
        fill_table_cell(t8.cell(i, 0), row_data[0])
        fill_table_cell(t8.cell(i, 1), row_data[1])

    # Table 9 changes
    t9 = tables[9]
    change_rows = [
        ("2026-05-22", "创建文档并完成引言、项目概述初稿", "项目启动", "项目组长", "开发团队"),
        ("2026-05-25", "补充业务需求分析与核心业务场景", "需求细化", "项目组长", "开发团队"),
        ("2026-05-28", "完善系统功能需求与用例分析", "功能设计", "项目组长", "开发团队"),
        ("2026-05-30", "完善非功能需求、优先级与术语表", "文档补充", "项目组长", "开发团队"),
        ("2026-06-04", "依据 README 与现有项目结构完成报告全面修订", "文档完善", "项目组长", "Codex/开发团队"),
    ]
    while len(t9.rows) < len(change_rows) + 1:
        t9.add_row()
    while len(t9.rows) > len(change_rows) + 1:
        t9._tbl.remove(t9.rows[-1]._tr)
    for i, row_data in enumerate(change_rows, start=1):
        for j, value in enumerate(row_data):
            fill_table_cell(t9.cell(i, j), value)

    doc.save(str(DOC_PATH))
    print(f"[OK] Updated {DOC_PATH}")


if __name__ == "__main__":
    main()
